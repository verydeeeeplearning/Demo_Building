from __future__ import annotations

import json
import math
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_JSON = ROOT / "docs" / "e2e_word_assets" / "e2e_12_case_results.json"
ASSET_DIR = ROOT / "docs" / "e2e_word_assets"
IMAGE_DIR = ASSET_DIR / "images"
OUT_DOCX = ROOT / "docs" / "deepagents_e2e_12_case_simulation_test.docx"

PAGE_WIDTH_DXA = 9360
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C4D2"
TEXT = "1F2937"
MUTED = "667085"
SUCCESS = "1F7A4D"


def east_asia_font(run, font_name="Malgun Gothic"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Deepagents E2E 12-case simulation test")
    east_asia_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 112, 133)


def add_title(doc: Document, data):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Deepagents E2E 12개 회로 시뮬레이션 테스트")
    east_asia_font(run)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    subtitle = (
        "학생의 탐색 대화부터 최종 입력, 내부 context/deepagents trace, "
        "deterministic validation/render/simulation 결과와 최종 회로 이미지를 한 문서에 묶은 E2E 실행 가이드입니다."
    )
    run = p.add_run(subtitle)
    east_asia_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)

    meta = doc.add_table(rows=3, cols=2)
    set_table_width(meta, [2200, PAGE_WIDTH_DXA - 2200])
    rows = [
        ("생성 시각", data["generatedAt"]),
        ("증거 생성", "buildContextPacket -> validateCircuitSpec -> compileRenderPlan -> compileSimulationPlan -> buildRunnableReport -> buildSolverGateResult"),
        ("범위", "WP-01부터 WP-12까지 각 1개 대표 학생 흐름, 총 12개 runnable simulation 케이스")
    ]
    for idx, (label, value) in enumerate(rows):
        meta.cell(idx, 0).text = label
        meta.cell(idx, 1).text = value
        set_cell_shading(meta.cell(idx, 0), LIGHT_BLUE)
        for cell in meta.row_cells(idx):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    east_asia_font(run)
                    run.font.size = Pt(9.5)
            if idx == 0:
                set_cell_shading(cell, LIGHT_BLUE if cell is meta.cell(idx, 0) else "FFFFFF")


def add_summary_table(doc: Document, cases):
    doc.add_heading("1. 전체 케이스 요약", level=1)
    table = doc.add_table(rows=1, cols=6)
    set_table_width(table, [850, 2450, 1850, 1900, 1300, 1010])
    headers = ["Case", "학생 목표", "Route", "Capability", "Gate", "Paths"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
    for case in cases:
        row = table.add_row()
        row.cells[0].text = f"{case['id'].split('-')[1]}\n{case['wp']}"
        row.cells[1].text = case["titleKo"]
        row.cells[2].text = case["context"]["routeId"]
        row.cells[3].text = ", ".join(case["context"]["capabilityIds"])
        row.cells[4].text = f"{case['validationReport']['status']} / {case['simulationPlan']['status']} / {case['runnableReport']['status']}"
        row.cells[5].text = str(len(case["simulationPlan"]["currentPaths"]))
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
            set_cell_border(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    east_asia_font(run)
                    run.font.size = Pt(8.5)


def find_font(size=18, bold=False):
    candidates = [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size, index=0)
        except Exception:
            continue
    return ImageFont.load_default()


def color_for_signal(signal: str):
    return {
        "power": (239, 68, 68),
        "ground": (31, 41, 55),
        "gpio": (37, 99, 235),
        "digital": (37, 99, 235),
        "button": (124, 58, 237),
        "analog": (14, 165, 233),
        "pwm": (249, 115, 22),
        "i2c-data": (37, 99, 235),
        "i2c-clock": (234, 179, 8),
        "uart": (20, 184, 166),
        "spi-clock": (234, 179, 8),
        "spi-data": (37, 99, 235),
        "spi-select": (124, 58, 237),
        "switched-ground": (75, 85, 99),
    }.get(signal, (90, 112, 140))


def project(point, bounds, width, height, margin=70):
    min_x, max_x, min_z, max_z = bounds
    scale_x = (width - margin * 2) / max(0.1, max_x - min_x)
    scale_z = (height - margin * 2 - 90) / max(0.1, max_z - min_z)
    scale = min(scale_x, scale_z)
    x = margin + (point["x"] - min_x) * scale
    y = margin + (point["z"] - min_z) * scale
    return int(x), int(y)


def draw_diagram(case, image_path: Path):
    render = case["renderPlan"]
    sim = case["simulationPlan"]
    endpoints = render.get("layout", {}).get("endpoints", {})
    parts = render["parts"]

    points = []
    for part in parts:
        pos = part["position"]
        footprint = part.get("footprint") or {}
        half_width = footprint.get("width", 0.45) / 2
        half_depth = footprint.get("depth", 0.35) / 2
        points.extend([
            {"x": pos["x"] - half_width, "z": pos["z"] - half_depth},
            {"x": pos["x"] + half_width, "z": pos["z"] + half_depth},
            pos,
        ])
    points.extend(endpoints.values())
    min_x = min(p["x"] for p in points) - 0.5
    max_x = max(p["x"] for p in points) + 0.5
    min_z = min(p["z"] for p in points) - 0.5
    max_z = max(p["z"] for p in points) + 0.5
    bounds = (min_x, max_x, min_z, max_z)

    width, height = 1400, 880
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = find_font(18)
    small = find_font(15)
    tiny = find_font(13)

    draw.rounded_rectangle((24, 24, width - 24, height - 24), radius=18, outline=(210, 220, 232), width=2, fill=(250, 252, 255))
    title = f"{case['wp']}  {case['titleKo']}"
    draw.text((48, 38), title, fill=(31, 77, 120), font=find_font(24))
    draw.text((48, 72), f"route={case['context']['routeId']}  status={case['validationReport']['status']}/{sim['status']}/{case['runnableReport']['status']}", fill=(71, 84, 103), font=small)

    for connection in render["connections"]:
        start = endpoints.get(f"{connection['from']['partId']}:{connection['from']['pin']}")
        end = endpoints.get(f"{connection['to']['partId']}:{connection['to']['pin']}")
        if not start or not end:
            continue
        x1, y1 = project(start, bounds, width, height)
        x2, y2 = project(end, bounds, width, height)
        y1 += 40
        y2 += 40
        color = color_for_signal(connection["signal"])
        draw.line((x1, y1, x2, y2), fill=color, width=5)
        mid = ((x1 + x2) // 2, (y1 + y2) // 2)
        draw.ellipse((mid[0] - 4, mid[1] - 4, mid[0] + 4, mid[1] + 4), fill=color)

    for part in parts:
        pos = part["position"]
        x, y = project(pos, bounds, width, height)
        y += 40
        footprint = part.get("footprint") or {}
        w = int(max(70, footprint.get("width", 0.45) * 150))
        h = int(max(46, footprint.get("depth", 0.35) * 150))
        fill_hex = footprint.get("visualStyle", {}).get("color", "#E5E7EB").lstrip("#")
        try:
            fill = tuple(int(fill_hex[i:i + 2], 16) for i in (0, 2, 4))
        except Exception:
            fill = (229, 231, 235)
        draw.rounded_rectangle((x - w // 2, y - h // 2, x + w // 2, y + h // 2), radius=12, fill=fill, outline=(55, 65, 81), width=2)
        label = f"{part.get('designator') or ''} {part['label']}".strip()
        if len(label) > 30:
            label = label[:27] + "..."
        tw = draw.textlength(label, font=tiny)
        draw.rounded_rectangle((x - tw / 2 - 8, y + h // 2 + 6, x + tw / 2 + 8, y + h // 2 + 26), radius=6, fill=(255, 255, 255), outline=(226, 232, 240))
        draw.text((x - tw / 2, y + h // 2 + 8), label, fill=(17, 24, 39), font=tiny)

    legend_y = height - 145
    draw.rounded_rectangle((48, legend_y, width - 48, height - 46), radius=14, fill=(244, 246, 249), outline=(216, 226, 237))
    draw.text((70, legend_y + 16), "Simulation evidence", fill=(31, 77, 120), font=font)
    paths = sim["currentPaths"]
    path_text = ", ".join(path["id"] for path in paths[:5]) if paths else "state-only context; no load-current path required"
    if len(paths) > 5:
        path_text += f" (+{len(paths) - 5})"
    draw.text((70, legend_y + 46), f"Current/signal paths: {path_text}", fill=(31, 41, 55), font=small)
    states = "; ".join(f"{state['componentId']}={state['state']}" for state in sim["expectedStates"][:3])
    draw.text((70, legend_y + 74), f"Expected states: {states or 'context state only'}", fill=(31, 41, 55), font=small)

    image.save(image_path)


def add_label_detail_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1900, PAGE_WIDTH_DXA - 1900])
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    east_asia_font(run)
                    run.font.size = Pt(9)
    return table


def add_log_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FAFC")
    set_cell_border(cell, color="D8E0EA")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Menlo"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_case(doc, case, image_path: Path):
    number = case["id"].split("-")[1]
    doc.add_heading(f"Case {number}. {case['titleKo']}", level=2)
    p = doc.add_paragraph()
    run = p.add_run(f"{case['wp']} 목적: ")
    east_asia_font(run)
    run.bold = True
    p.add_run(case["purpose"])
    for run in p.runs:
        east_asia_font(run)

    doc.add_heading("학생 탐색 흐름", level=3)
    for idx, turn in enumerate(case["studentTurns"], start=1):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(f"{idx}. 학생: ")
        east_asia_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        run = p.add_run(turn)
        east_asia_font(run)

    add_label_detail_table(doc, [
        ("최종 입력", case["finalPrompt"]),
        ("기대 판정", case["expectedAssertion"]),
        ("Route", case["context"]["routeId"]),
        ("Capability", ", ".join(case["context"]["capabilityIds"])),
        ("Prompt budget", case["context"]["promptBudget"]),
    ])

    doc.add_heading("내부 Deepagents Trace", level=3)
    add_log_block(doc, case["deepagentLog"])

    doc.add_heading("Gate Summary", level=3)
    sim_paths = ", ".join(path["id"] for path in case["simulationPlan"]["currentPaths"]) or "state-only"
    warnings = case["validationReport"]["warnings"] + case["simulationPlan"]["warnings"]
    add_label_detail_table(doc, [
        ("Validation", f"{case['validationReport']['status']} | errors={len(case['validationReport']['errors'])} | warnings={len(case['validationReport']['warnings'])}"),
        ("Render", f"parts={len(case['renderPlan']['parts'])} | wires={len(case['renderPlan']['connections'])} | warnings={len(case['renderPlan']['warnings'])}"),
        ("Simulation", f"{case['simulationPlan']['status']} | paths={sim_paths}"),
        ("Runnable", f"{case['runnableReport']['status']} | runnable={case['runnableReport']['runnable']}"),
        ("Solver Gate", solver_gate_summary(case)),
        ("주요 warning", " / ".join(warnings[:2]) if warnings else "none"),
    ])

    doc.add_heading("최종 Simulation 회로", level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = caption.add_run(f"Figure {number}. {case['titleKo']} simulation circuit evidence")
    east_asia_font(cap)
    cap.italic = True
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor(102, 112, 133)


def add_acceptance(doc, cases):
    doc.add_heading("3. 실행 전 Acceptance Checklist", level=1)
    checks = [
        "각 케이스에서 학생의 마지막 입력이 기존 blocked artifact가 아니라 새 synthesis request로 처리된다.",
        "Deepagents trace에 route, selected bundle, candidate parts, validation, render, simulation, runnable gate, solver gate가 모두 남는다.",
        "Validation이 valid이고 simulation이 valid일 때만 최종 simulation 회로 이미지를 승인한다.",
        "통신/logic/interface 케이스는 실제 네트워크, 페어링, 정밀 측정, exact timing을 주장하지 않는다.",
        "지원되는 저전압 회로는 build confirmation 이후 render canvas와 current/signal overlay로 이어져야 한다.",
        "문서에 기재된 prompt budget은 maxPromptChars 이하로 유지되어야 한다."
    ]
    for check in checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        run = p.add_run("□ ")
        east_asia_font(run)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        run = p.add_run(check)
        east_asia_font(run)

    doc.add_heading("4. 12개 케이스 최종 판정", level=1)
    table = doc.add_table(rows=1, cols=5)
    set_table_width(table, [1000, 3350, 1500, 1700, 1810])
    for idx, header in enumerate(["Case", "목표", "Validation", "Simulation", "Runnable"]):
        table.cell(0, idx).text = header
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
    for case in cases:
        row = table.add_row()
        row.cells[0].text = f"{case['id'].split('-')[1]}\n{case['wp']}"
        row.cells[1].text = case["titleKo"]
        row.cells[2].text = case["validationReport"]["status"]
        row.cells[3].text = case["simulationPlan"]["status"]
        row.cells[4].text = case["runnableReport"]["status"]
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    east_asia_font(run)
                    run.font.size = Pt(9)


def solver_gate_summary(case):
    gate = case.get("solverGateResult")
    if not gate:
        return "not recorded"
    return (
        f"{gate.get('mode')} | visible={gate.get('visibleSimulation')} | "
        f"buildReady={gate.get('buildReady')} | activity={gate.get('simulationActivity')}"
    )


def main():
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(ASSET_JSON.read_text(encoding="utf-8"))
    cases = data["cases"]

    image_paths = {}
    for case in cases:
        image_path = IMAGE_DIR / f"{case['id']}.png"
        draw_diagram(case, image_path)
        image_paths[case["id"]] = image_path

    doc = Document()
    style_document(doc)
    add_title(doc, data)
    add_summary_table(doc, cases)

    doc.add_heading("2. 케이스별 E2E 흐름", level=1)
    lead = doc.add_paragraph(
        "각 케이스는 학생이 회로 목표를 탐색하고, 마지막 입력으로 회로를 확정한 뒤, "
        "내부 context/deepagents tool chain이 어떤 판정을 내렸는지 보여줍니다."
    )
    for run in lead.runs:
        east_asia_font(run)

    for index, case in enumerate(cases):
        if index:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_case(doc, case, image_paths[case["id"]])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_acceptance(doc, cases)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
