from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "hardware_object_drag_and_adjustment_evidence.docx"
EVIDENCE = ROOT / "docs" / "evidence" / "hardware-adjustment"


def main() -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc)
    add_decision_callout(doc)
    add_summary_table(doc)
    add_adjustment_matrix(doc)
    add_image_section(doc)
    add_verification_table(doc)
    add_qa_notes(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    run = title.add_run("Hardware Object Drag and Adjustment Evidence")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} KST").italic = True
    subtitle.paragraph_format.space_after = Pt(10)


def add_decision_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Gate interpretation correction. ")
    r.bold = True
    p.add_run(
        "The visual-move screenshot is not proof that a hardware placement gate accepted the moved layout. "
        "It is a dirty visual preview: canonical CircuitSpec and verified wire routes remain unchanged, "
        "and Run/current controls stay disabled until reset or hardware placement resolution. "
        "Gate-pass evidence is the hardware-move resolver output plus the adjustment-mode E2E matrix."
    )


def add_summary_table(doc: Document) -> None:
    doc.add_heading("Implementation Evidence", level=1)
    rows = [
        ("Phase 1", "Rule gate as automatic presentation adjustment", "Four E2E adjustment-mode cases remain visible without unsupported controls.", "Passed"),
        ("Phase 2", "Scene graph and debug snapshot", "Part, wire, label, and preview-wire groups expose deterministic debug metadata.", "Passed"),
        ("Phase 3", "Visual-only movement", "Visual transforms move rendered objects and dashed preview wires without mutating CircuitSpec.", "Passed"),
        ("Phase 4", "Hardware movement resolver", "Dragged component is sent to placement API, snapped, rerouted, and returned as canonical artifacts.", "Passed"),
        ("QA fix", "Library-only parts excluded from default PCB stage", "Demo circuit no longer shows non-circuit library parts as if they were verified hardware.", "Passed"),
    ]
    table = add_table(doc, ["Scope", "What Changed", "Evidence", "Status"], rows, [1150, 2350, 4500, 1360])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def add_adjustment_matrix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Adjustment Mode Matrix", level=1)
    rows = [
        ("Diagnostic", "Safe 3D diagnostic scene remains visible while build/current claims stay scoped.", "Run/current disabled unless verified paths exist.", "Phase 1 E2E"),
        ("Placeholder", "Generic registered placeholder may render when exact footprint evidence is missing.", "Exact geometry and pin claims disabled.", "Phase 1 E2E"),
        ("Safe Equivalent", "Unsafe original is replaced with a safe displayed equivalent.", "Original never build-ready; equivalent controls follow its own validation.", "Phase 1 E2E"),
        ("State Only", "Pin-map, layout, and static state context can render without fake current flow.", "Current animation off; static inspection on.", "Phase 1 E2E"),
    ]
    add_table(doc, ["Mode", "Visible Adjustment", "Control Policy", "Proof"], rows, [1600, 3750, 2700, 1310])


def add_image_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Browser Evidence Screenshots", level=1)
    images = [
        ("01-demo-baseline.png", "Baseline verified OLED demo scene. Library-only preview parts are no longer rendered in the PCB stage."),
        ("02-visual-move-preview.png", "Visual-only movement preview. This is intentionally not a hardware gate pass; the scene is marked as visual arrangement and Run/current controls are disabled."),
        ("03-hardware-move-resolved.png", "Hardware placement resolution. The dragged LED layout is returned through the placement API as a canonical adjusted render plan."),
    ]
    for index, (filename, caption) in enumerate(images):
        if index:
            doc.add_page_break()
        path = EVIDENCE / filename
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.paragraph_format.keep_with_next = True
        run = caption_paragraph.add_run(caption)
        run.bold = True
        doc.add_picture(str(path), width=Inches(5.25))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_verification_table(doc: Document) -> None:
    doc.add_heading("Verification Commands", level=1)
    rows = [
        ("Syntax", "node --check src/main.js && node --check src/stageScene.js", "Passed"),
        ("Typecheck", "npm run typecheck", "Passed"),
        ("Unit", "npm run test:unit", "318 passed"),
        ("Build", "npm run build", "Passed; existing Vite large chunk warning only"),
        ("Targeted E2E", "Phase 1 adjustment matrix, 3D render, visual move, hardware move", "7 passed"),
        ("Focused post-QA E2E", "3D render, visual move, hardware move after libraryOnly fix", "3 passed"),
        ("Evidence capture", "node scripts/generate_hardware_evidence.mjs", "Passed"),
    ]
    add_table(doc, ["Gate", "Command or Case", "Result"], rows, [1450, 5600, 2310])


def add_qa_notes(doc: Document) -> None:
    doc.add_heading("QA Notes", level=1)
    notes = [
        "The image that showed an off-board blue part was not a hardware gate-pass result; it was a visual preview and also exposed a libraryOnly rendering problem.",
        "Default PCB rendering now excludes libraryOnly parts so the visible stage matches the verified circuit under inspection.",
        "Visual movement is still allowed for inspection, but it stays dirty and disables Run/current controls until reset.",
        "Hardware movement is the canonical path: pointer-up sends a deterministic placement intent, then the server returns adjusted artifacts.",
        "The post-QA E2E test now asserts that hardware move changes wireRouteHash and remains buildReady after resolver output.",
    ]
    for note in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(note)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        shade_cell(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
    doc.add_paragraph()
    return table


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, width in enumerate(widths):
            tc_w = row.cells[index]._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.type = "dxa"
            tc_w.w = width


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


if __name__ == "__main__":
    main()
